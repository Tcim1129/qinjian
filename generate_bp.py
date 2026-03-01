from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_cell_text(cell, text, bold=False, font_size=10.5):
    """创建单元格文本"""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(font_size)
    run.bold = bold
    return paragraph


# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(12)

# 添加封面
doc.add_paragraph()
doc.add_paragraph()

# 主标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("亲健")
run.font.name = "黑体"
run.font.size = Pt(48)
run.bold = True

doc.add_paragraph()

# 副标题
subtitle1 = doc.add_paragraph()
subtitle1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle1.add_run("——青年亲密关系AI健康管理平台")
run.font.name = "黑体"
run.font.size = Pt(22)

doc.add_paragraph()

# 商业计划书
subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run("商业计划书")
run.font.name = "黑体"
run.font.size = Pt(24)

subtitle3 = doc.add_paragraph()
subtitle3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle3.add_run("（省赛一等奖优化版）")
run.font.name = "楷体"
run.font.size = Pt(16)

# 添加空行
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# 日期
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("2025年2月")
run.font.name = "宋体"
run.font.size = Pt(14)

# 分页
doc.add_page_break()

# ================== 第一章 ==================
# 标题
h1 = doc.add_heading("第一章 执行摘要", level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h1.runs:
    run.font.name = "黑体"
    run.font.size = Pt(18)
    run.bold = True

# 1.1 项目定位
h2 = doc.add_heading("1.1 项目定位", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

h3 = doc.add_heading("1.1.1 核心定义", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run("亲健项目以").font.name = "宋体"
p.add_run('"亲密关系健康管理"').font.name = "宋体"
p.add_run(
    "为核心定位，开创性地将人工智能技术与权威心理学理论深度融合，构建面向18-40岁青年群体的专业化服务平台。这一市场定位的精准性体现在三个维度："
).font.name = "宋体"

# 三点内容
bullet_points = [
    ("场景聚焦", "——深耕情侣、夫妻、挚友三类核心人际关系，形成差异化竞争壁垒；"),
    ("技术驱动", "——通过数据驱动的预防性干预，实现关系健康的主动管理；"),
    (
        "平台化设计",
        "——为后续B端服务拓展预留充足空间，使C端用户积累转化为企业级服务的核心资产。",
    ),
]

for bold_text, normal_text in bullet_points:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(bold_text)
    run.bold = True
    run.font.name = "宋体"
    p.add_run(normal_text).font.name = "宋体"

# 1.1.2 目标用户
h3 = doc.add_heading("1.1.2 目标用户", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run(
    "亲健的目标用户群体覆盖18-40岁青年，基于深刻的用户行为洞察进行精细化分层运营："
).font.name = "宋体"

age_groups = [
    (
        "18-25岁高校学生群体",
        "处于亲密关系探索期，对轻量化、游戏化的关系管理工具接受度高，适合以免费试用+社交裂变快速获取种子用户；",
    ),
    (
        "26-35岁年轻职场人群",
        "面临工作压力与亲密关系维护的双重挑战，对高效、专业的关系改善方案付费意愿更强，适合以专业报告+专家咨询提升客单价；",
    ),
    (
        "36-40岁已婚人群",
        "进入关系深度经营阶段，对夫妻沟通、育儿协作等场景的专业指导需求迫切，适合以家庭账户+育儿专题增强用户粘性。",
    ),
]

for bold_text, normal_text in age_groups:
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    run.font.name = "宋体"
    p.add_run(normal_text).font.name = "宋体"

# 1.1.3 核心价值
h3 = doc.add_heading("1.1.3 核心价值", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run("亲健的核心价值主张建立在").font.name = "宋体"
p.add_run('"科学评估"与"个性化改善"').font.name = "宋体"
p.add_run(
    "两大支柱之上。科学性体现在：引入约翰·戈特曼（John Gottman）的亲密关系研究理论作为底层框架；采用双视角数据采集机制，打破传统单一方自我报告的局限性。个性化体现在：AI系统基于双视角数据为每对关系生成定制化的改善任务推送，形成从认知唤醒到行为改变的完整干预链条。"
).font.name = "宋体"

# 1.2 核心理念
h2 = doc.add_heading("1.2 核心理念", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

h3 = doc.add_heading("1.2.1 AI增强真实关系", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run(
    '2025年上半年全球AI陪伴应用收入达到8200万美元，但头部产品存在结构性困境：Character.AI付费率不足1%，Replika虽盈利却深陷"人与虚拟关系"的伦理争议。这些产品的共同特征是将AI定位为"关系替代者"，用户与虚拟角色建立情感连接，反而可能削弱真实人际关系的能力和质量。'
).font.name = "宋体"

p = doc.add_paragraph()
p.add_run("亲健选择差异化路径——").font.name = "宋体"
p.add_run('"AI增强真实关系"').font.name = "宋体"
p.add_run(
    "，将技术能力聚焦于真实人际关系的诊断与改善支持。双视角数据验证机制确保AI的介入始终以促进真实互动为目标；隐私沙盒架构从技术上杜绝了原始数据的单方获取可能；改善任务系统推送的具体行动均需双方线下完成。"
).font.name = "宋体"

# 表1：隐私沙盒三层架构
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表1：隐私沙盒三层架构").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=4)
table.style = "Table Grid"

headers = ["层级", "功能定位", "数据内容", "访问权限"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "Layer 1",
        "原始数据层：本地存储，不上传",
        "语音原始音频、详细文字记录、消费记录",
        "仅用户本人",
    ],
    [
        "Layer 2",
        "AI分析层：仅AI可见特征向量",
        "语音情感特征向量、行为模式标签、关系健康评分",
        "仅AI系统",
    ],
    [
        "Layer 3",
        "洞察层：双方可见抽象洞察",
        "情绪同步度75%、本周深度交流3次、改善建议",
        "关系双方",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 1.2.2 解决高期待-低能力的关系焦虑
h3 = doc.add_heading("1.2.2 解决高期待-低能力的关系焦虑", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run("当代青年群体面临的").font.name = "宋体"
p.add_run('"高期待-低能力"困境').font.name = "宋体"
p.add_run(
    "是亲健产品设计的核心问题导向。从社会文化维度来看，浪漫爱情理想的媒体传播与亲密关系教育的社会缺位形成张力；从代际比较维度来看，当代青年对情感满足、个人成长、平等对话等关系质量维度的要求显著提高，但实际的关系经营技能却未能同步发展。"
).font.name = "宋体"

# 1.3 核心创新点
h2 = doc.add_heading("1.3 核心创新点", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

h3 = doc.add_heading("1.3.1 双视角数据验证", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run("双视角数据验证是亲健").font.name = "宋体"
p.add_run("最具技术壁垒的核心创新").font.name = "宋体"
p.add_run(
    '，其创新性体现在数据采集方法论的根本变革。传统的关系评估工具依赖单一主体自我报告，存在显著的主观偏差；亲健的双视角机制要求关系双方独立、同步完成每日打卡，AI系统对双方数据进行交叉比对，计算"情绪同步度"（双方情绪评分差的绝对值）、"互动平衡度"（双方主动性评分的比值）等衍生指标，识别关系投入的不对称性与互动模式。'
).font.name = "宋体"

# 表2：心理学理论底座
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表2：心理学理论底座与产品应用").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=5, cols=3)
table.style = "Table Grid"

headers = ["理论来源", "核心贡献", "产品化应用"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        '戈特曼"四骑士"理论',
        "批评、蔑视、防御、stonewalling预测关系危机",
        "AI风险识别特征与温和预警机制",
    ],
    [
        '戈特曼"5:1积极比例"',
        "稳定关系需5次积极互动抵消1次消极互动",
        '"互动平衡度"核心指标设计',
    ],
    ["依恋类型理论", "安全型/焦虑型/回避型/恐惧型依恋模式", "差异化改善建议推送"],
    ["情绪聚焦疗法（EFT）", "情绪调节与关系改变技术", "改善任务内容设计"],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 1.4 发展阶段与目标
h2 = doc.add_heading("1.4 发展阶段与目标", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表3：亲健四阶段发展目标与财务预测").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=5, cols=7)
table.style = "Table Grid"

headers = ["阶段", "时间", "注册关系", "付费用户", "月收入", "净利润", "关键里程碑"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "校内验证期",
        "1-2月",
        "50-100对",
        "4-8人",
        "¥100-200",
        "-¥22,000",
        "产品价值验证",
    ],
    [
        "校内推广期",
        "3-6月",
        "400-500对",
        "32-40人",
        "¥800-1,000",
        "-¥30,000",
        "单校模型验证",
    ],
    ["B端启动期", "6-12月", "3,000对", "240人", "¥12,000+", "¥100,000", "整体盈亏平衡"],
    [
        "规模化期",
        "1-2年",
        "50,000对",
        "4,000人",
        "¥200,000+",
        "¥1,000,000+",
        "全国市场覆盖",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 分页
doc.add_page_break()

# ================== 第二章 ==================
h1 = doc.add_heading("第二章 市场分析", level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h1.runs:
    run.font.name = "黑体"
    run.font.size = Pt(18)

h2 = doc.add_heading("2.1 宏观市场环境", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

h3 = doc.add_heading("2.1.1 情绪经济崛起", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run(
    "情绪经济作为新兴消费形态，正在重塑中国消费市场的结构特征。根据市场研究数据，"
).font.name = "宋体"
p.add_run(
    "2026年中国情绪消费市场规模预计达到2.72万亿元，年复合增长率超过15%"
).font.name = "宋体"
p.add_run("。").font.name = "宋体"
p.add_run('37.6%的年轻消费者将"价值共鸣"作为购买决策的首要因素').font.name = "宋体"
p.add_run("，显著高于价格敏感度（21.3%）与品牌知名度（18.7%）。").font.name = "宋体"

h3 = doc.add_heading("2.1.2 心理健康服务高速增长", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run("中国心理健康服务市场正处于高速成长期，").font.name = "宋体"
p.add_run(
    "2025年市场规模预计突破500亿元，年增长率超过25%，但渗透率仍不足5%"
).font.name = "宋体"
p.add_run(
    "，与发达国家15%-20%的水平存在显著差距。核心瓶颈在于：供给端，专业心理咨询师不足10万人，且分布高度集中于一线城市；需求端，社会stigma（病耻感）导致大量潜在用户不愿寻求专业帮助。"
).font.name = "宋体"

h3 = doc.add_heading("2.1.3 AI情感陪伴赛道爆发", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run("AI情感陪伴应用在2025年迎来爆发式增长，").font.name = "宋体"
p.add_run("上半年全球收入达到8200万美元，预计全年突破1.2亿美元").font.name = "宋体"
p.add_run(
    '。但头部产品的结构性问题为后来者提供了差异化空间：Character.AI付费率不足1%，商业模式高度依赖广告；Replika虽实现盈利，但"人与虚拟关系"的伦理争议不断。'
).font.name = "宋体"

# 2.2 亲密关系困境与需求
h2 = doc.add_heading("2.2 亲密关系困境与需求", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

h3 = doc.add_heading("2.2.1 婚恋危机", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run("中国正经历深刻的婚恋结构转型。").font.name = "宋体"
p.add_run(
    "2023年离婚率超过40%，较2010年上升近3倍；结婚率从2013年的9.9‰骤降至2023年的4.4‰，创历史新低"
).font.name = "宋体"
p.add_run(
    '。这一趋势的背后是"高期待-低能力"的普遍落差——当代青年对亲密关系的质量期待空前提高，但实际的关系经营能力却未同步提升。'
).font.name = "宋体"

h3 = doc.add_heading("2.2.2 友谊淡化", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run("社交媒体时代的人际关系呈现出").font.name = "宋体"
p.add_run('"连接增加、深度减少"的悖论特征').font.name = "宋体"
p.add_run("。").font.name = "宋体"
p.add_run('80%的大学生表示有亲密朋友但"感觉关系在淡化"').font.name = "宋体"
p.add_run(
    '——"点赞之交"取代深度对话，碎片化即时通讯降低了主动联系的动机，地理分散与时间挤压使友谊维护更加困难。'
).font.name = "宋体"

h3 = doc.add_heading("2.2.3 预调研验证", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run("亲健团队在本校开展的50人预调研为产品设计提供了实证基础：").font.name = "宋体"

# 表4：预调研
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表4：亲健预调研核心发现").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=5, cols=3)
table.style = "Table Grid"

headers = ["调研问题", "结果", "产品启示"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    ["恋爱/友谊中遇到过沟通问题？", '82%"经常"或"有时"', "沟通辅助是核心痛点"],
    ["希望有工具帮助改善关系？", '72%"愿意尝试"', "需求真实存在"],
    ["能接受每天2分钟记录？", '85%"可以接受"', "轻量化设计可行"],
    ["愿意为关系改善付费？", '40%"小额愿意"', "付费意愿需培养"],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 2.3 校内市场验证场景
h2 = doc.add_heading("2.3 校内市场验证场景", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

h3 = doc.add_heading("2.3.1 市场容量测算", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

# 表5：市场容量
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表5：校内市场容量测算").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=5, cols=4)
table.style = "Table Grid"

headers = ["人群细分", "规模", "占比", "特征分析"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    ["在校生总数", "15,000人", "100%", "—"],
    ["情侣（含暧昧期）", "4,500人", "30%", "主打人群：痛点感知强、互动频率高"],
    ["有亲密朋友者", "12,000人", "80%", "测试人群：市场教育成本低"],
    ["可触达率", "~70%", "—", "3,150对潜在用户"],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 分页
doc.add_page_break()

# ================== 第三章 ==================
h1 = doc.add_heading("第三章 竞品分析", level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h1.runs:
    run.font.name = "黑体"
    run.font.size = Pt(18)

h2 = doc.add_heading("3.1 直接竞品：情侣社交赛道", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表6：直接竞品
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表6：直接竞品对比分析").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=4)
table.style = "Table Grid"

headers = ["产品", "用户规模", "核心问题", "亲健差异化优势"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "小恩爱",
        "3000万",
        "功能堆砌（47个二级功能），缺乏AI分析，7日留存降至28%",
        "极简设计+AI洞察+价值导向",
    ],
    [
        "Between",
        "2500万",
        "本土化不足，ARPU仅$0.5/月，已退出中国市场主流",
        "深度本土化+技术-政策-文化三位一体",
    ],
    [
        "微爱",
        "5000万下载",
        "定位追踪引发隐私争议，评分跌至3.2分",
        "健康导向而非监控导向",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 3.2 间接竞品
h2 = doc.add_heading("3.2 间接竞品：心理健康与AI陪伴赛道", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表7：间接竞品
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表7：间接竞品对比分析").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=4)
table.style = "Table Grid"

headers = ["产品", "模式", "核心问题", "亲健差异化路径"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "Soul",
        "陌生人社交",
        "关系悖论：匹配成功即用户流失，年流失率>60%",
        '服务"经营关系"而非"寻找关系"',
    ],
    ["Replika", "AI虚拟陪伴", "无法解决真实关系困境，人机关系脆弱", "服务真实人际关系"],
    [
        "壹心理",
        "专业心理咨询",
        "单次300-800元，决策周期长，转化率低",
        "轻量化+预防性+日常化",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 3.3 功能矩阵对比
h2 = doc.add_heading("3.3 功能矩阵对比", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表8：功能矩阵
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表8：功能矩阵全面对比").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=7, cols=8)
table.style = "Table Grid"

headers = ["功能维度", "小恩爱", "Between", "Soul", "Replika", "壹心理", "亲健", "差异"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    if i == 6:
        set_cell_shading(cell, "E8F5E9")
    else:
        set_cell_shading(cell, "D5E8F0")

data_rows = [
    ["情侣关系支持", "✅", "✅", "❌", "❌", "❌", "✅", "—"],
    ["夫妻关系支持", "❌", "❌", "❌", "❌", "❌", "✅", "独有"],
    ["挚友关系支持", "❌", "❌", "❌", "❌", "❌", "✅", "独有"],
    ["双视角互动记录", "❌", "❌", "❌", "❌", "❌", "✅", "独有"],
    ["AI关系健康分析", "❌", "❌", "❌", "✅", "❌", "✅", "真实关系"],
    ["隐私沙盒保护", "❌", "❌", "❌", "❌", "❌", "✅", "独有"],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        cell = table.rows[i].cells[j]
        create_cell_text(cell, text)
        if j == 6:
            set_cell_shading(cell, "E8F5E9")

# 分页
doc.add_page_break()

# ================== 第四章 ==================
h1 = doc.add_heading("第四章 产品方案", level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h1.runs:
    run.font.name = "黑体"
    run.font.size = Pt(18)

h2 = doc.add_heading("4.1 核心功能架构", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

h3 = doc.add_heading("4.1.1 关系类型选择系统", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run("亲健的用户首次使用流程经过精心设计，三步完成onboarding：").font.name = "宋体"

steps = [
    (
        "Step 1：选择关系类型",
        "。三种选项对应差异化设计：👫 情侣（恋爱中）侧重互动频率、深度交流、情感表达；💑 夫妻（已婚）侧重家务分工、育儿协作、二人世界时间；👭 挚友（闺蜜/兄弟）侧重主动联系、深度分享、共同体验。",
    ),
    (
        "Step 2：邀请对方绑定",
        "。生成6位数字绑定码，对方输入即完成关系确认，避免复杂的账号注册与权限设置。",
    ),
    (
        "Step 3：开始每日打卡",
        "。晚21:00推送提醒，2分钟内完成，将行为门槛控制在可接受范围内。",
    ),
]

for bold_text, normal_text in steps:
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    run.font.name = "宋体"
    p.add_run(normal_text).font.name = "宋体"

h3 = doc.add_heading("4.1.2 双视角/单机兼容打卡", level=3)
for run in h3.runs:
    run.font.name = "黑体"
    run.font.size = Pt(12)

# 表9：打卡流程
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表9：双视角打卡流程设计").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=5, cols=4)
table.style = "Table Grid"

headers = ["步骤", "问题设计", "数据类型", "分析价值"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "Step 1",
        "今天心情如何？😊😐😢😠",
        "情绪评分（1-4分）",
        "个人情绪基线，关系情绪传染检测",
    ],
    ["Step 2", "今天和TA互动多吗？", "互动频率+主动性", "互动平衡度，关系投入差异识别"],
    [
        "Step 3",
        "今天有深度交流吗？（30分钟以上）",
        "深度交流布尔值",
        "关系质量核心指标，危机预警信号",
    ],
    ["Step 4", "今日任务完成了吗？", "任务完成率", "行为改变追踪，干预效果评估"],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 4.2 隐私沙盒架构
h2 = doc.add_heading("4.2 隐私沙盒架构", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表10：隐私沙盒详细
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表10：隐私沙盒三层架构详解").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=4)
table.style = "Table Grid"

headers = ["层级", "核心机制", "技术实现", "用户价值"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "Layer 1",
        "原始数据层：本地存储，不上传",
        "设备端加密存储，任何网络传输均不发生",
        "绝对隐私保障，消除数据泄露风险",
    ],
    [
        "Layer 2",
        "AI分析层：仅AI可见特征向量",
        "端侧MiniCPM-V提取特征，差分隐私添加噪声（ε=1-8）",
        "AI有效分析，但任何人无法还原原始内容",
    ],
    [
        "Layer 3",
        "洞察层：双方可见抽象洞察",
        "云端Qwen2.5-VL推理，输出聚合指标与建议",
        "有价值但不可逆向，促进共同行动",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 4.3 AI关系健康报告
h2 = doc.add_heading("4.3 AI关系健康报告", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表11：订阅产品
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表11：分层订阅产品设计").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=5, cols=4)
table.style = "Table Grid"

headers = ["产品层级", "定价", "核心内容", "目标用户"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "每日简报（免费）",
        "¥0",
        "情绪同步度、互动平衡度、关系温度趋势、风险信号、明日建议",
        "所有注册用户，培养使用习惯",
    ],
    [
        "周报深度版",
        "¥9.9/周",
        "7天趋势图、关系阶段判断、同龄对比、改善方案",
        "轻度用户，首次付费转化",
    ],
    [
        "月报专业版",
        "¥29.9/月",
        "90天趋势、个性化计划、AI语音解读、专家咨询入口",
        "中度用户，高价值服务",
    ],
    [
        "年度会员",
        "¥199/年",
        "全功能解锁、专属客服、线下活动邀请",
        "高粘性核心用户，提升LTV",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 4.4 游戏化改善任务
h2 = doc.add_heading("4.4 游戏化改善任务", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

p = doc.add_paragraph()
p.add_run('"关系树"成长体系').font.name = "宋体"
run = p.runs[0]
run.bold = True
p.add_run(
    "是亲健游戏化机制的核心载体。用户完成打卡和任务获得积分，推动关系树从种子→萌芽→幼苗→小树→大树→森林的成长，视觉化进度反馈强化行为动机。"
).font.name = "宋体"

# 表12：任务库
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表12：差异化任务库设计").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=5)
table.style = "Table Grid"

headers = ["关系类型", "任务示例", "积分", "完成验证", "难度"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    ["情侣", "睡前说晚安、每周一次约会、写一封情书", "5-30分", "双方确认", "初级-高级"],
    [
        "夫妻",
        "分担一项家务、安排二人世界时间、感谢对方",
        "5-30分",
        "双方确认",
        "初级-高级",
    ],
    [
        "挚友",
        "每周主动联系一次、分享一个秘密、一起运动",
        "5-30分",
        "单方确认",
        "初级-高级",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 分页
doc.add_page_break()

# ================== 第五章 ==================
h1 = doc.add_heading("第五章 技术架构与数据壁垒", level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h1.runs:
    run.font.name = "黑体"
    run.font.size = Pt(18)

h2 = doc.add_heading("5.1 大模型选型策略", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表13：大模型选型
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表13：端-云协同架构设计").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=3, cols=5)
table.style = "Table Grid"

headers = ["部署位置", "模型", "参数", "核心优势", "应用场景"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "端侧",
        "MiniCPM-V 2.6",
        "2.8B",
        "手机NPU可运行，数据不出设备",
        "语音情感分析、文本特征提取",
    ],
    [
        "云端",
        "Qwen2.5-VL-7B",
        "7B",
        "中文场景深度优化，复杂推理能力强",
        "关系趋势分析、个性化建议生成",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 5.2 微调策略
h2 = doc.add_heading("5.2 微调策略", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

p = doc.add_paragraph()
p.add_run(
    "采用QLoRA轻量化微调，核心参数：可训练参数~0.5%，显存需求8-12GB，训练时间1-2天，单张RTX 4090即可完成。这一技术选择使初创团队无需大规模算力投入即可实现模型定制化。"
).font.name = "宋体"

# 5.3 核心护城河
h2 = doc.add_heading("5.3 核心护城河——私有数据库", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

p = doc.add_paragraph()
p.add_run("亲健明确将").font.name = "宋体"
p.add_run('"私有数据库"而非"专利"').font.name = "宋体"
p.add_run(
    '定位为核心护城河。通过种子用户积累国内首个"双视角关系行为特征数据库"，其独特性体现在：双视角结构（双方独立记录，交叉验证）、时间序列深度（持续追踪关系发展全过程）、多模态特征（语音、文本、行为多维度）、中文语境（本土文化表达的情感模式）。'
).font.name = "宋体"

# 5.4 数据安全与合规
h2 = doc.add_heading("5.4 数据安全与合规", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表14：合规框架
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表14：合规框架与技术实现").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=3)
table.style = "Table Grid"

headers = ["法规", "核心要求", "亲健应对措施"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    ["《个人信息保护法》", "敏感信息需单独同意", "分层授权设计，一键撤回功能"],
    ["《数据安全法》", "数据分类分级管理", "三级数据分类体系（原始/特征/洞察）"],
    [
        "《生成式人工智能服务管理暂行办法》",
        "算法备案、安全评估",
        '上线前完成备案，输出标注"AI生成"',
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 分页
doc.add_page_break()

# ================== 第六章 ==================
h1 = doc.add_heading("第六章 商业模式与盈利预测", level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h1.runs:
    run.font.name = "黑体"
    run.font.size = Pt(18)

h2 = doc.add_heading("6.1 收入结构", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表15：收入结构
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表15：三层次收入结构设计").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=4)
table.style = "Table Grid"

headers = ["收入类型", "目标占比", "核心产品", "定价策略"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "C端订阅",
        "40%",
        "周报¥9.9、月报¥29.9、年费¥199",
        "价值阶梯设计，免费→低价→高价转化",
    ],
    [
        "增值服务",
        "20%",
        "专家咨询¥50-200/次、情感课程¥99/门",
        "平台撮合+质量管控，抽成20-30%",
    ],
    [
        "B端服务",
        "40%",
        "高校科研合作¥5-10万/年、企业EAP¥20-100/人/年",
        "科研合作建立背书，EAP服务规模化变现",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 6.2 成本结构
h2 = doc.add_heading("6.2 成本结构", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表16：成本结构
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表16：成本结构精细化测算").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=5, cols=3)
table.style = "Table Grid"

headers = ["成本类型", "金额/单价", "说明"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    ["固定成本", "¥454/年", "阿里云服务器¥99/年+小程序认证¥300/年"],
    ["变动成本-单次AI报告", "¥0.001-0.0065", "基础版vs深度版token消耗差异"],
    ["变动成本-月成本（1000用户）", "¥50-200", "日活60%，50%生成周报，20%生成月报"],
    ["推广成本-CAC目标", "¥50-80", "朋友圈裂变¥5-10、班级群渗透¥10-15、宿舍地推¥15-20"],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 6.3 分阶段盈利预测
h2 = doc.add_heading("6.3 分阶段盈利预测", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表17：财务预测
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表17：三阶段财务预测与里程碑").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=4)
table.style = "Table Grid"

headers = ["阶段", "关键指标", "财务结果", "核心目标"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "校内验证期（1-2月）",
        "50对注册，2-3%付费转化",
        "净利润-¥22,000",
        "验证产品价值，获取用户行为数据",
    ],
    [
        "校内推广期（3-6月）",
        "400对注册，LTV/CAC=2.5",
        "6个月累计-¥30,000",
        "单校模型验证，优化转化漏斗",
    ],
    [
        "B端启动期（6-12月）",
        "3,000对，3所高校+5家企业",
        "年度净利润¥100,000",
        "整体盈亏平衡，B端收入启动",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 分页
doc.add_page_break()

# ================== 第七章 ==================
h1 = doc.add_heading("第七章 运营计划：14天冷启动实验", level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h1.runs:
    run.font.name = "黑体"
    run.font.size = Pt(18)

h2 = doc.add_heading("7.1 MVP实验", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

p = doc.add_paragraph()
p.add_run(
    '样本：30对真实情侣，为期14天手动打卡实验。验证指标："关系指数波动曲线"（每日关系健康评分的变化趋势）、"AI建议后的行为改变率"（用户实际执行改善任务的比例）。'
).font.name = "宋体"

h2 = doc.add_heading("7.2 渠道组合", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表18：渠道组合
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表18：零成本/低成本渠道组合").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=5)
table.style = "Table Grid"

headers = ["渠道", "具体动作", "预算", "目标", "效率特征"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "朋友圈裂变",
        "转发产品海报，邀请成功双方各获7天周报免费体验",
        "¥0",
        "15对",
        "社交信任度高，病毒传播潜力大",
    ],
    [
        "班级群渗透",
        '联系班级干部，以"班级情感建设"名义推广',
        "¥0",
        "15对",
        "转化率稳定，组织动员效应",
    ],
    [
        "宿舍地推",
        "周末晚间扫楼，发放定制书签（带绑定码）",
        "¥200",
        "20对",
        "转化率高，人效低，适合初期验证",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

h2 = doc.add_heading("7.3 B端启动策略", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表19：高校合作
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表19：高校合作渐进路径").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=5, cols=4)
table.style = "Table Grid"

headers = ["层级", "合作内容", "亲健收益", "关键成功因素"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    ["Level 1", "免费科研合作", "种子用户+学术信任", "产品专业性与研究价值证明"],
    ["Level 2", "数据合作", "学术背书+联合论文", "数据质量与合规保障"],
    ["Level 3", "联合发表", "品牌权威+招投标资质", "研究成果的学术影响力"],
    ["Level 4", "区域独家", "¥5-10万/年+竞争壁垒", "前期合作积累的信任与依赖"],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

h2 = doc.add_heading("7.4 关键运营指标", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表20：关键指标
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表20：关键运营指标体系").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=5, cols=5)
table.style = "Table Grid"

headers = ["维度", "指标", "目标值", "监控频率", "战略意义"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "北极星指标",
        "健康关系数（持续打卡7天以上）",
        "阶段增长",
        "每周",
        "反映产品核心价值的实际交付",
    ],
    ["增长", "绑定成功率", ">70%", "每日", "双边产品激活的关键门槛"],
    ["活跃", "日活率", ">50%", "每日", "用户习惯养成的基础"],
    ["商业", "付费转化率", "2-3%", "每周", "商业模式可行性的关键验证"],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 分页
doc.add_page_break()

# ================== 第八章 ==================
h1 = doc.add_heading("第八章 团队与组织架构", level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h1.runs:
    run.font.name = "黑体"
    run.font.size = Pt(18)

h2 = doc.add_heading("8.1 核心团队配置", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表21：团队配置
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表21：核心团队配置与分工").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=5, cols=4)
table.style = "Table Grid"

headers = ["角色", "股权占比", "职责范围", "关键任务"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    ["CEO/产品", "40%", "产品战略、融资对接", "B端合作谈判、融资对接、产品方向决策"],
    ["CTO/AI", "25%", "技术架构、AI集成", "隐私计算实现、模型优化、技术团队建设"],
    ["CMO/运营", "20%", "品牌建设、用户增长", "校内推广、内容营销、用户社区运营"],
    ["COO/B端", "15%", "商务拓展、资源整合", "高校/企业合作、资源整合、运营效率优化"],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

h2 = doc.add_heading("8.2 外部资源网络", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表22：外部资源
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表22：外部资源网络").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=4)
table.style = "Table Grid"

headers = ["资源类型", "合作对象", "合作内容", "当前进度"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    ["专业背书", "校心理中心", "联合研究、工具授权、用户推荐", "洽谈中"],
    ["技术基础设施", "阿里云", "云服务器、AI大模型API、创业扶持", "已开通"],
    ["服务供给", "认证情感咨询师", "在线咨询服务、课程内容", "初步接触"],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 分页
doc.add_page_break()

# ================== 第九章 ==================
h1 = doc.add_heading("第九章 风险分析与应对", level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h1.runs:
    run.font.name = "黑体"
    run.font.size = Pt(18)

h2 = doc.add_heading("9.1 产品风险", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表23：产品风险
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表23：产品风险与应对").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=3, cols=3)
table.style = "Table Grid"

headers = ["风险", "具体表现", "应对措施"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "用户隐私顾虑",
        "担心情感数据泄露，不敢真实记录",
        "隐私沙盒架构，原始数据本地处理，双方只看抽象洞察",
    ],
    [
        "伴侣关系破裂",
        "解绑行为可能引发冲突或负面情绪",
        "解绑需双方确认或7天冷静期，解绑后推送关怀内容",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

h2 = doc.add_heading("9.2 商业风险", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表24：商业风险
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表24：商业风险与应对").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=3, cols=3)
table.style = "Table Grid"

headers = ["风险", "具体表现", "应对措施"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "LTV<CAC",
        "用户获取成本高于生命周期价值",
        "B端收入补贴，多渠道获客，严控CAC在¥50-80",
    ],
    [
        "付费意愿不足",
        "免费用户转化付费的比例低于预期",
        "价值阶梯设计，免费试用→低价周报→年费路径",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

h2 = doc.add_heading("9.3 竞争风险", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表25：竞争风险
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表25：竞争风险与应对").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=3)
table.style = "Table Grid"

headers = ["风险", "具体表现", "应对措施"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "大厂入局",
        "互联网巨头推出类似产品",
        "速度窗口（6-12月建立数据壁垒）+核心架构专利保护",
    ],
    [
        "竞品抄袭功能",
        "竞争对手复制双视角或隐私沙盒",
        "双视角交叉验证算法与情感隐私沙盒机制已申请专利",
    ],
    [
        "技术迭代",
        "大模型技术快速演进，现有方案落后",
        "持续优化AI层，数据壁垒+心理学量表独家授权",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 分页
doc.add_page_break()

# ================== 第十章 ==================
h1 = doc.add_heading("第十章 融资规划", level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h1.runs:
    run.font.name = "黑体"
    run.font.size = Pt(18)

h2 = doc.add_heading("10.1 资金需求与用途", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表26：融资规划
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表26：三阶段融资规划").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=6)
table.style = "Table Grid"

headers = ["轮次", "时间", "金额", "资金来源", "核心用途", "里程碑"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    ["种子轮", "当前", "¥5,000", "自筹+校内创业基金", "MVP开发、校内验证", "50-100对"],
    [
        "天使轮",
        "6-12月",
        "¥50-100万",
        "天使投资人/校友基金",
        "5校扩张、团队建设",
        "单校模型",
    ],
    [
        "Pre-A轮",
        "12-18月",
        "¥300-500万",
        "VC机构",
        "全国市场覆盖、B端深度探索",
        "10万用户",
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

h2 = doc.add_heading("10.2 投资人沟通策略", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表27：投资人策略
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表27：投资人沟通策略").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=3)
table.style = "Table Grid"

headers = ["投资人类型", "核心关注点", "亲健的故事"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    [
        "财务投资人（VC）",
        "市场规模、增长潜力",
        '"情绪经济中的关系健康入口"——2.72万亿情绪消费市场中的垂直赛道领导者',
    ],
    [
        "战略投资人（大厂）",
        "场景协同、生态价值",
        '"关系数据的战略价值"——亲密关系数据对社交、内容、消费场景的赋能潜力',
    ],
    [
        "产业投资人",
        "专业壁垒、社会价值",
        '"数字化心理健康服务创新"——AI技术降低服务门槛，提升国民关系健康水平',
    ],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 分页
doc.add_page_break()

# ================== 第十一章 ==================
h1 = doc.add_heading("第十一章 附录", level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h1.runs:
    run.font.name = "黑体"
    run.font.size = Pt(18)

h2 = doc.add_heading("11.1 核心指标目标汇总", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表28：核心指标汇总
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表28：核心指标目标汇总").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=4, cols=7)
table.style = "Table Grid"

headers = ["阶段", "时间", "注册关系", "付费用户", "月收入", "净利润", "关键里程碑"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    ["校内验证", "1-2月", "50对", "4人", "¥100", "-¥22,000", "产品价值验证"],
    ["校内推广", "3-6月", "400对", "32人", "¥800", "-¥20,000", "单校模型验证"],
    ["B端启动", "6-12月", "3,000对", "240人", "¥12,000+", "¥100,000", "整体盈亏平衡"],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

h2 = doc.add_heading("11.2 支持材料清单", level=2)
for run in h2.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)

# 表29：材料清单
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("表29：支持材料清单").font.name = "宋体"
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

table = doc.add_table(rows=6, cols=3)
table.style = "Table Grid"

headers = ["材料类型", "具体内容", "当前状态"]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    create_cell_text(cell, header, bold=True)
    set_cell_shading(cell, "D5E8F0")

data_rows = [
    ["用户研究", "50人预调研原始数据与洞察报告", "已归档"],
    ["竞品分析", "小恩爱、Between、Soul、Replika、壹心理深度研究", "已完成"],
    ["产品设计", "产品原型（Figma）与交互流程图", "进行中"],
    ["技术文档", "系统架构设计、数据库Schema、API接口规范", "初稿完成"],
    ["财务模型", "LTV/CAC测算、敏感性分析、融资规划Excel", "已完成"],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        create_cell_text(table.rows[i].cells[j], text)

# 结尾声明
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "本商业计划书基于2024-2025年最新市场数据和行业研究编制，融合隐私沙盒架构、温和AI洞察、保守财务预测、B端双轮驱动四大优化，致力于成为中国青年亲密关系健康管理的首选平台。"
)
run.font.name = "宋体"
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

# 保存文档
doc.save("亲健商业计划书.docx")
print("文档已生成: 亲健商业计划书.docx")
