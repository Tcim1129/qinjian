from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Users\colour\Desktop\亲健")
OUT = ROOT / "亲健_三创比赛提交版_规范版.docx"


def set_run_font(run, size=12, bold=False, name="Arial", east="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style = "Heading 1" if level == 1 else "Heading 2"
    p.style = doc.styles[style]
    r = p.add_run(text)
    set_run_font(r, 16 if level == 1 else 13, True, east="黑体")
    return p


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.28)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r, 12)
    return p


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "TOC " + '\\o "1-3" ' + "\\h \\z \\u"
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    txt = paragraph.add_run("目录将在 Word 中更新。右键目录并选择“更新域”。")
    set_run_font(txt, 11)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 10)


def add_table(doc, caption, headers, rows):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    return table


def add_image(doc, image_name, caption, width=5.3):
    path = ROOT / "screenshots" / image_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_footer_page_numbers(doc):
    for sec in doc.sections:
        footer = sec.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run("第 ")
        set_run_font(r1, 10)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "1"
        r.append(t)
        fld.append(r)
        p._p.append(fld)
        r2 = p.add_run(" 页")
        set_run_font(r2, 10)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    for _ in range(3):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("亲健——青年亲密关系 AI 健康管理平台")
    set_run_font(r, 20, True, east="黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("第十六届全国大学生电子商务“创新、创意及创业”挑战赛项目报告书")
    set_run_font(r, 15, True, east="黑体")

    for _ in range(4):
        doc.add_paragraph("")

    cover_lines = [
        "项目类别：常规赛",
        "参赛方向：社会服务 / 电子商务",
        "项目负责人：吴秀秀",
        "团队成员：黄菁蓝、叶笙尧、钟昊桐、郑梓滢",
        "指导教师：闻敏、黄军",
        "所属学院：[请填写学院名称]",
        "完成时间：[请填写日期]",
    ]
    for line in cover_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, 12)

    add_page_break(doc)

    add_heading(doc, "摘要", 1)
    abstract_paragraphs = [
        "亲健是一款面向青年亲密关系场景的 AI 健康管理平台，聚焦情侣、夫妻、挚友等真实关系中的沟通不足、情绪表达失衡、异地互动困难、冲突修复能力弱、关系状态难以持续追踪等问题，致力于通过数字化工具与人工智能能力，为用户提供低门槛、可持续、可视化的关系健康管理服务。",
        "项目围绕“AI 增强真实关系”这一核心理念展开，构建了“每日打卡—关系分析—AI 报告—改善建议—长期追踪”的产品闭环。系统当前重点建设内容包括：用户认证与登录、关系配对、双视角打卡、AI 日报/周报/月报、关系树成长机制、危机预警、异地关系模块、依恋分析、关系健康测试、里程碑管理、课程与专家咨询入口等。平台现已形成“后端服务 + Web 工作台 + 微信小程序 + App 端”四部分协同推进的系统结构，项目正从概念方案阶段转向产品打磨与功能验证阶段。",
        "从项目建设看，后端已基于 FastAPI 完成认证、配对、打卡、报告、关系树、危机预警、任务、异地活动、里程碑、社群提示、智能陪伴等模块接口；Web 端已形成关系健康工作台；微信小程序已搭建首页、打卡、发现、报告、我的及多个扩展页面；App 端基于 uni-app 已实现首页、打卡、报告、发现、个人中心等核心页面并持续联调优化。当前项目已具备较为完整的系统骨架和比赛展示基础。",
        "关键词：亲密关系、AI 健康管理、微信小程序、Web 工作台、App、多端协同、关系报告、危机预警、社会服务、电子商务",
    ]
    for text in abstract_paragraphs:
        add_body_paragraph(doc, text)

    add_page_break(doc)

    add_heading(doc, "目录", 1)
    add_toc(doc.add_paragraph())

    add_page_break(doc)

    sections = [
        (
            "一、项目概述",
            [
                (
                    "1.1 项目背景",
                    "当代青年在亲密关系中面临越来越明显的现实挑战。很多年轻人并未建立起稳定的关系经营能力，导致沟通不足、冲突升级、情绪累积、联系变淡等问题频繁出现。",
                ),
                (
                    "1.2 项目定位",
                    "亲健定位为面向青年真实亲密关系场景的 AI 健康管理平台，强调真实关系、持续记录和辅助分析三项原则。",
                ),
                (
                    "1.3 项目目标",
                    "项目希望建设围绕真实关系经营的数字化平台，构建支持 Web、小程序、App 的多端产品形态，并形成“打卡—分析—报告—改善—追踪”的核心闭环。",
                ),
                (
                    "1.4 项目当前状态",
                    "截至目前，项目已搭建后端 API 服务、Web 关系健康工作台、微信小程序主路径页面及 App 端核心页面，并已实现认证、配对、打卡、报告、关系树、危机预警等基础模块。",
                ),
            ],
        ),
        (
            "二、市场分析",
            [
                (
                    "2.1 用户需求分析",
                    "青年群体普遍存在记录关系状态、查看关系趋势、提前识别问题和获得可执行改善建议的真实需求，同时高度关注隐私可控和移动端便捷使用。",
                ),
                (
                    "2.2 目标用户群体",
                    "项目当前主要面向情侣用户、夫妻用户、挚友用户和单人体验用户，覆盖真实关系中的多种典型场景。",
                ),
                (
                    "2.3 同类产品分析",
                    "同类产品主要包括情侣互动工具、心理咨询平台和 AI 情感陪伴类产品。相较之下，亲健的切入点是围绕真实关系长期管理构建多端轻量工具。",
                ),
                (
                    "2.4 项目切入点",
                    "项目聚焦真实关系、持续记录、关系改善和多端协同，试图形成差异化的产品路径。",
                ),
            ],
        ),
        (
            "三、产品与服务",
            [
                (
                    "3.1 产品核心理念",
                    "亲健强调“AI 增强真实关系，而不是替代真实关系”，目标是帮助用户看见关系状态、理解互动节奏并获得改善支持。",
                ),
                (
                    "3.2 产品主路径",
                    "产品主路径为：登录注册 → 关系配对 → 每日打卡 → AI 分析 → 报告输出 → 改善建议 → 长期追踪。",
                ),
                (
                    "3.3 核心功能模块",
                    "当前核心功能包括用户认证、关系配对、每日打卡、AI 报告、关系树成长机制、危机预警及扩展服务入口。",
                ),
                (
                    "3.4 服务延伸方向",
                    "在基础功能之外，项目还预留了关系课程、专家咨询、校园活动合作和青年服务场景合作等延伸方向。",
                ),
            ],
        ),
        (
            "四、项目技术实现与系统介绍",
            [
                (
                    "4.1 系统总体架构",
                    "本项目采用“多前端 + 统一后端服务”的系统架构模式，由 Web、小程序、App 端共同接入后端业务服务。",
                ),
                (
                    "4.2 后端实现情况",
                    "后端基于 FastAPI，已形成 auth、pairs、checkins、reports、tree、crisis、tasks、longdistance、milestones、community、agent 等模块化结构。",
                ),
                (
                    "4.3 Web 端实现情况",
                    "Web 端当前定位为关系健康工作台，已支持登录、首页、打卡、发现、报告和个人中心等主要页面。",
                ),
                (
                    "4.4 微信小程序实现情况",
                    "微信小程序端已搭建首页、打卡、发现、报告、我的及多个扩展功能页面，适合关系打卡和轻量使用场景。",
                ),
                (
                    "4.5 App 端实现情况",
                    "App 端基于 uni-app 构建，已完成首页、打卡、报告、发现、个人中心等核心页面骨架。",
                ),
            ],
        ),
        (
            "五、营销策略与推广路径",
            [
                (
                    "5.1 推广总体思路",
                    "当前阶段推广策略更适合采用“轻量验证、场景传播、逐步扩展”的路径。",
                ),
                (
                    "5.2 初期推广路径",
                    "初期重点围绕校园冷启动、关系测试引流和节日场景传播展开。",
                ),
            ],
        ),
        (
            "六、商业模式与财务规划",
            [
                (
                    "6.1 商业模式概述",
                    "项目当前商业模式主要包括 C 端订阅、增值服务和 B 端合作三层结构。",
                ),
                (
                    "6.2 当前阶段商业化原则",
                    "当前阶段应坚持先验证产品价值，再逐步推动变现的原则。",
                ),
                (
                    "6.3 阶段性财务规划思路",
                    "财务规划建议围绕产品验证期、场景试点期和合作拓展期三个阶段展开。",
                ),
            ],
        ),
        (
            "七、风险管理与边界说明",
            [
                (
                    "7.1 项目边界",
                    "亲健定位为关系健康辅助工具，而非心理治疗系统或医疗诊断系统。",
                ),
                (
                    "7.2 项目主要风险",
                    "主要风险包括隐私风险、AI 输出风险、多端联调风险和产品定位风险。",
                ),
            ],
        ),
        (
            "八、团队介绍与实施保障",
            [
                (
                    "8.1 团队情况说明",
                    "当前项目在技术实现层面主要由核心成员独立推进完成，团队其他成员围绕文档、调研、展示和支持协同推进。",
                ),
                (
                    "8.2 团队分工",
                    "建议在提交版中真实呈现团队职责，突出技术开发主要由一人承担的实际情况。",
                ),
            ],
        ),
        (
            "九、发展规划",
            [
                (
                    "9.1 短期规划",
                    "短期围绕主流程打通、多端联调、体验优化和材料补充推进。",
                ),
                (
                    "9.2 中期规划",
                    "中期将强化异地关系、智能陪伴、课程与专家内容以及用户反馈闭环。",
                ),
                (
                    "9.3 长期规划",
                    "长期将继续围绕青年关系健康管理方向，形成更稳定的多端数字健康平台。",
                ),
            ],
        ),
        (
            "十、结论",
            [
                (
                    "10.1 总结",
                    "亲健项目已形成较为完整的系统结构和产品雏形，具备较好的三创比赛项目基础与后续扩展空间。",
                ),
            ],
        ),
    ]

    for sec_title, subs in sections:
        add_heading(doc, sec_title, 1)
        for sub_title, text in subs:
            add_heading(doc, sub_title, 2)
            add_body_paragraph(doc, text)

    add_heading(doc, "规范化图表", 2)

    add_table(
        doc,
        "表 1-1 项目当前建设内容一览",
        ["建设层级", "当前状态", "说明"],
        [
            ("后端服务", "已完成核心骨架", "已具备认证、配对、打卡、报告等模块"),
            ("Web 端", "已完成主要页面", "可用于演示与桌面端体验"),
            ("微信小程序", "已完成主路径页面", "首页、打卡、发现、报告、我的等已搭建"),
            ("App 端", "已完成核心页面", "基于 uni-app，正在持续联调优化"),
            ("部署能力", "已具备基础条件", "已有 Docker Compose、Nginx、云部署配置"),
        ],
    )

    add_table(
        doc,
        "表 2-1 目标用户群体细分",
        ["用户类型", "典型场景", "核心需求"],
        [
            (
                "情侣用户",
                "日常互动、沟通不足、异地维护",
                "打卡、趋势分析、异地辅助、改善建议",
            ),
            ("夫妻用户", "沟通协调、长期陪伴、冲突修复", "长期追踪、报告、危机提醒"),
            ("挚友用户", "联系减少、深度交流不足", "关系维持、提醒、里程碑记录"),
            ("单人体验用户", "未绑定对象、先体验产品", "自我记录、情绪整理、功能预览"),
        ],
    )

    add_table(
        doc,
        "表 3-1 当前核心功能模块一览",
        ["模块", "主要作用", "当前状态"],
        [
            ("登录认证", "建立账户与身份识别", "已实现"),
            ("关系配对", "建立双人关系空间", "已实现"),
            ("每日打卡", "收集关系状态数据", "已实现"),
            ("AI 报告", "形成趋势分析与建议", "已实现"),
            ("关系树", "增强成长反馈", "已实现"),
            ("危机预警", "识别异常状态", "已实现"),
            ("异地关系", "支持异地互动场景", "已接入"),
            ("依恋测试", "帮助理解互动模式", "已接入"),
            ("课程/专家/会员", "服务化延展入口", "已搭建入口"),
        ],
    )

    add_table(
        doc,
        "表 4-1 后端模块与实现情况",
        ["模块", "主要能力"],
        [
            ("auth", "注册、登录、手机号登录、微信登录、修改资料、修改密码"),
            ("pairs", "创建配对、加入配对、关系摘要、解绑、昵称备注"),
            ("checkins", "打卡提交、今日打卡、历史打卡、连续天数"),
            ("reports", "生成日报、周报、月报，获取最新报告、历史报告、趋势"),
            ("upload", "图片上传、语音上传"),
            ("tree", "关系树状态、浇水成长"),
            ("crisis", "危机状态、危机记录、预警处理、资源推荐"),
            ("tasks", "每日任务、任务完成、依恋分析"),
            ("longdistance", "异地活动创建、完成、异地健康指数"),
            ("milestones", "里程碑创建、查看、回顾生成"),
            ("community", "经营提示、通知列表、已读处理"),
            ("agent", "智能陪伴会话、聊天记录、对话能力"),
        ],
    )

    add_table(
        doc,
        "表 5-1 当前适合的推广路径",
        ["推广方式", "主要场景", "目标"],
        [
            ("校园冷启动", "班级群、宿舍、熟人链", "获取种子用户"),
            ("关系测试引流", "小程序测试入口", "降低首次体验门槛"),
            ("节日传播", "情人节、七夕、毕业季", "提升话题传播度"),
            ("校园活动合作", "社团、心理活动周", "提升真实触达"),
        ],
    )

    add_table(
        doc,
        "表 6-1 商业模式结构",
        ["层级", "形式", "主要内容"],
        [
            ("C 端订阅", "周报、月报、年度会员", "面向个人用户的持续服务"),
            ("增值服务", "专家咨询、课程、深度测试", "面向高需求用户的附加服务"),
            ("B 端合作", "高校、青年服务、员工关怀合作", "面向组织场景的拓展方向"),
        ],
    )

    add_table(
        doc,
        "表 6-2 阶段性财务与目标思路",
        ["阶段", "核心任务", "重点投入", "核心指标"],
        [
            (
                "产品验证期",
                "完成系统打磨与校园试用",
                "开发、部署、测试",
                "注册量、配对率、活跃率",
            ),
            (
                "场景试点期",
                "探索付费服务与试点合作",
                "内容、活动、用户支持",
                "报告使用率、复购率、付费率",
            ),
            (
                "合作拓展期",
                "探索高校及青年服务合作",
                "推广、服务、内容供给",
                "合作数量、留存率、服务稳定性",
            ),
        ],
    )

    add_table(
        doc,
        "表 7-1 主要风险与应对措施",
        ["风险类型", "具体表现", "应对措施"],
        [
            (
                "隐私风险",
                "涉及情绪、语音、关系状态等敏感信息",
                "强化权限控制与隐私说明",
            ),
            ("AI 输出风险", "报告或建议可能不准确", "增加辅助性质说明与边界提示"),
            ("多端联调风险", "终端体验和接口状态不一致", "优先打通主链路并分阶段联调"),
            ("产品定位风险", "功能过多导致主线分散", "围绕“打卡—报告—改善”持续收敛"),
        ],
    )

    add_table(
        doc,
        "表 8-1 团队分工表",
        ["角色", "姓名", "主要职责"],
        [
            ("项目负责人", "吴秀秀", "统筹项目方向、比赛材料、汇报展示、项目推进"),
            (
                "核心技术开发",
                "[请填写实际开发者姓名]",
                "后端、Web、小程序、App 的系统设计与开发",
            ),
            ("产品/文档协助", "钟昊桐", "功能整理、文档配合、材料支持"),
            ("运营/内容协助", "叶笙尧", "展示准备、传播思路、活动支持"),
            ("商务/调研协助", "郑梓滢", "市场调研、竞品整理、合作资料支持"),
        ],
    )

    add_page_break(doc)
    add_heading(doc, "附录 A 系统页面与功能展示说明", 1)
    appendix_images = [
        ("home.png", "图 A-1 首页"),
        ("checkin.png", "图 A-2 打卡页"),
        ("report.png", "图 A-3 报告页"),
        ("discover.png", "图 A-4 发现页"),
        ("profile.png", "图 A-5 个人中心页"),
        ("auth-login.png", "图 A-6 登录页"),
        ("auth-register.png", "图 A-7 注册页"),
        ("attachment-test.png", "图 A-8 依恋测试页"),
        ("health-test.png", "图 A-9 关系健康测试页"),
        ("longdistance.png", "图 A-10 异地关系页"),
        ("community.png", "图 A-11 社区页"),
        ("challenges.png", "图 A-12 挑战赛页"),
        ("courses.png", "图 A-13 课程页"),
        ("experts.png", "图 A-14 专家咨询页"),
        ("membership.png", "图 A-15 会员中心页"),
    ]
    for image_name, caption in appendix_images:
        add_image(doc, image_name, caption)

    add_heading(doc, "附录 B 建议补充的比赛材料", 1)
    for item in [
        "学院与团队信息完整页",
        "系统架构图",
        "产品功能流程图",
        "当前开发进度说明图",
        "关键接口或模块说明图",
        "实际测试记录或试用反馈",
        "演示视频二维码",
        "部署说明摘要",
        "竞品分析简表",
        "项目答辩讲解提纲",
    ]:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, 12)

    set_footer_page_numbers(doc)
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    build()
